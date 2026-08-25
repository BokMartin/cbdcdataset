import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "paper/current.docx"
PDF = ROOT / "paper/current.pdf"
TITLE = "INSTITUTIONAL DETERMINANTS OF PRIVACY DESIGN IN CENTRAL BANK DIGITAL CURRENCIES"
HEADER = "International Days of Statistics and Economics"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def statement_count(path: Path) -> int:
    total = 0
    with path.open(encoding="utf-8") as stream:
        for line in stream:
            record = json.loads(line)
            total += sum(len(unit["statements"]) for unit in record["units"])
    return total


def paragraph_after(document: Document, label: str):
    for index, paragraph in enumerate(document.paragraphs[:-1]):
        if paragraph.text.strip() == label:
            return document.paragraphs[index + 1]
    raise AssertionError(f"missing paragraph after {label}")


def reference_paragraphs(document: Document):
    start = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "References")
    end = next(i for i, paragraph in enumerate(document.paragraphs[start + 1 :], start + 1) if paragraph.text.strip() == "Contact")
    return [paragraph for paragraph in document.paragraphs[start + 1 : end] if paragraph.text.strip()]


def main() -> None:
    document = Document(DOCX)
    text = document_text(document)
    pages = PdfReader(PDF).pages
    summary = load_json(ROOT / "results/v10_2e_ensemble/analysis_summary.json")
    ensemble = load_json(ROOT / "results/v10_2e_ensemble/macro_results.json")["variants"]["ensemble"]
    package = load_json(ROOT / "validation/extraction_v10_2_exploratory/freeze/package_manifest.json")

    openai_statements = statement_count(
        ROOT / "validation/extraction_v10_2_exploratory/runs/2026-08-25_openai_production/final/openai_extractions_v10_2e_canonical.jsonl"
    )
    claude_statements = statement_count(
        ROOT / "validation/extraction_v10_2_exploratory/runs/2026-08-25_claude_production/claude_extractions_v10_2e_canonical.jsonl"
    )

    abstract_words = re.findall(r"\b[\w’.-]+\b", paragraph_after(document, "Abstract").text)
    references = reference_paragraphs(document)
    title = document.paragraphs[0]
    authors = document.paragraphs[1]
    body_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if len(paragraph.text.split()) >= 25
        and not paragraph.text.startswith(("DATA PENDING", "Source:"))
    ]
    zero_spacing = all(
        (paragraph.paragraph_format.space_before is None or paragraph.paragraph_format.space_before.pt == 0)
        and (paragraph.paragraph_format.space_after is None or paragraph.paragraph_format.space_after.pt == 0)
        for paragraph in body_paragraphs
    )
    with zipfile.ZipFile(DOCX) as archive:
        names = set(archive.namelist())
        xml = archive.read("word/document.xml")

    checks = {
        "uppercase exact title": title.text == TITLE and title.text == title.text.upper(),
        "title format": title.alignment == WD_ALIGN_PARAGRAPH.CENTER and all(run.bold and round(run.font.size.pt) == 16 for run in title.runs),
        "authors format": authors.alignment == WD_ALIGN_PARAGRAPH.CENTER and all(run.bold and round(run.font.size.pt) == 14 for run in authors.runs),
        "abstract 150-200 words": 150 <= len(abstract_words) <= 200,
        "five keywords": "central bank digital currency, privacy, content analysis, large language models, reproducibility" in text,
        "JEL classification": "JEL Code: E42, E58, C83" in text,
        "body paragraph spacing": zero_spacing,
        "page range 8-10": 8 <= len(pages) <= 10,
        "conference header every page": all(HEADER in (page.extract_text() or "") for page in pages),
        "two figures and two tables": len(document.inline_shapes) == 2 and len(document.tables) == 2,
        "fifteen references": len(references) == 15,
        "contact details": all(value in text for value in ["bokm00@vse.cz", "dominik.stroukal@mup.cz", "Prague University of Economics and Business", "Metropolitan University Prague"]),
        "no comments or content controls": "word/comments.xml" not in names and b"<w:sdt" not in xml,
        "corpus counts": package["counts"] == {"documents": 113, "pages": 3963, "units": 3963, "requests": 661, "render_pages": 345} and all(value in text for value in ["113", "3,963"]),
        "provider statement counts": openai_statements == 5675 and claude_statements == 5683 and "11,358" in text,
        "candidate population": summary["candidate_population"] == 6949 and "6,949" in text,
        "origin decomposition": summary["origin_type"] == {"both": 4307, "claude_only": 1327, "openai_only": 1315} and all(value in text for value in ["4,307", "1,315", "1,327"]),
        "agreement": round(100 * summary["agreement_among_both"]["exact_code_set_rate"], 1) == 41.5 and all(value in text for value in ["41.5%", "44.9%", "62.8%"]),
        "calibration boundary": all(value in text for value in ["precision 1.000", "recall 0.625", "0.977", "0.656", "0.980", "0.766", "0.633–0.881", "failed the frozen recall thresholds", "reserve was therefore not opened"]),
        "pending human boundary": summary["human_validation"] == {"candidate_population": 6949, "candidate_sample": 365, "dual_empty_population": 359, "dual_empty_sample": 36, "status": "pending_blind_double_review"} and all(value in text for value in ["DATA PENDING BEFORE SUBMISSION", "365 candidates", "36 of 359"]),
        "dominant centres": summary["dominant_centres"]["ensemble"] == {"cash_substitution": 4, "financial_inclusion": 4, "monetary_transmission": 6, "payment_modernization": 12, "sovereignty_competition": 7, "state_control": 14},
        "measurement relation": round(ensemble["measurement_relation"]["r"], 3) == 0.219 and "r = 0.219" in text,
        "macro headline": round(ensemble["posture"]["shadow"]["r"], 3) == -0.286 and round(ensemble["vocabulary"]["basel"]["r"], 3) == -0.292 and all(value in text for value in ["−0.286", "−0.292"]),
        "Holm conclusion": all(ensemble[measure][condition]["p_holm"] > 0.05 for measure in ["vocabulary", "posture"] for condition in ["vdem", "cbi", "shadow", "basel"]) and "No institutional correlate survives multiplicity correction" in text,
        "release placeholders": all(value in text for value in ["https://github.com/BokMartin/cbdc-msed-v10", "https://bokmartin.github.io/cbdc-msed-v10/", "TO BE ACTIVATED AND VERIFIED BEFORE SUBMISSION"]),
        "no stale v9 headline": "6,139" not in text and "5,624" not in text,
    }
    failed = [name for name, passed in checks.items() if not passed]
    if failed:
        raise AssertionError("paper checks failed: " + ", ".join(failed))
    print(f"paper: {len(checks)} checks passed; pages={len(pages)}; abstract_words={len(abstract_words)}; references={len(references)}")


if __name__ == "__main__":
    main()
