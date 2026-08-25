#!/usr/bin/env python3
"""Replace the superseded 15% coding statement in the current manuscript."""

from pathlib import Path

from docx import Document


OLD = (
    "Inter-coder reliability for the multi-dimensional coding scheme is being assessed with "
    "Krippendorff’s α on a 15% double-coded sample (Krippendorff, 2019), which is appropriate "
    "because the coding spans several dimensions and is not restricted to two coders."
)
NEW = (
    "Candidate validity and inter-coder reliability are being assessed on a pre-frozen blind "
    "probability sample of 365 of 6,949 candidates (5.25%; 95% worst-case margin ±4.99 percentage "
    "points), independently coded by both authors before disagreement-only consensus. We will "
    "report design-weighted acceptance, percent agreement, Krippendorff’s α, and Gwet’s AC1; "
    "retrieval recall remains the separately frozen calibration estimate (Krippendorff, 2019)."
)
OLD_REVISION = (
    "Revised version: 19 August 2026 (statistical audit, measurement update and DCash stage "
    "correction; human validation in progress; full audit trail at cbdcdataset.org)."
)
NEW_REVISION = (
    "Revised version: 25 August 2026 (statistical audit, measurement update, DCash stage "
    "correction, and pre-frozen sampled human-validation design; validation in progress; "
    "full audit trail at cbdcdataset.org)."
)


def main() -> None:
    path = Path(__file__).resolve().parents[1] / "paper" / "current.docx"
    document = Document(path)
    method_matches = []
    revision_matches = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if OLD in run.text:
                method_matches.append(run)
            if OLD_REVISION in run.text:
                revision_matches.append(run)
    if len(method_matches) > 1 or len(revision_matches) > 1:
        raise RuntimeError("duplicate manuscript targets")
    if method_matches:
        method_matches[0].text = method_matches[0].text.replace(OLD, NEW)
    elif not any(NEW in paragraph.text for paragraph in document.paragraphs):
        raise RuntimeError("sampled-validation methods text not found")
    if revision_matches:
        revision_matches[0].text = revision_matches[0].text.replace(OLD_REVISION, NEW_REVISION)
    elif not any(NEW_REVISION in paragraph.text for paragraph in document.paragraphs):
        raise RuntimeError("revision statement not found")
    document.save(path)


if __name__ == "__main__":
    main()
